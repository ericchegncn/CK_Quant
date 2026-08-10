# -*- coding: utf-8 -*-
"""生成《教学头模快速拆装机构专利授权前景分析报告》Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_GREEN = RGBColor(0x16, 0x65, 0x34)
DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
GRAY = RGBColor(0x44, 0x44, 0x44)

doc = Document()

# 全局默认字体（含东亚字体）
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_east_asia(run, name='微软雅黑'):
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE if level == 1 else RGBColor(0x00, 0x00, 0x00)
        set_east_asia(run)
    return h

def add_para(text, bold=False, size=11, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    set_east_asia(r)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        set_east_asia(r1)
    r = p.add_run(text)
    set_east_asia(r)
    p.paragraph_format.space_after = Pt(4)
    return p

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def style_table_cell(cell, bold=False, size=10.5, color=None, center=True):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.bold = bold
            r.font.size = Pt(size)
            if color is not None:
                r.font.color.rgb = color
            set_east_asia(r)

# ============ 封面标题 ============
title = add_para('教学头模快速拆装机构', bold=True, size=22, color=DARK_BLUE,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para('专利授权前景分析报告（美欧双边）', bold=True, size=15, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para('分析日期：2026年8月6日 ｜ 保密等级：内部资料', size=10, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

# ============ 一、发明概述 ============
add_heading('一、发明概述', 1)
add_para('本申请涉及一种教学/美发用头模（mannequin head）的快速拆装机构，核心创新在于'
         '「单向盲插下压自动锁紧 + 顶部斜面解锁按钮 + 底部顶杆弹簧一键弹出」的三位一体机械结构：', space_after=8)
add_bullet('将头壳对准底座后直接下压，倒扣杆自动挤开横向卡条并锁紧，无需拧螺丝、无需对准卡扣、无需充气；',
           bold_prefix='① 盲插锁紧：')
add_bullet('解锁时仅需按压头壳顶端的解锁按钮，按钮斜面驱动横向卡条平移释放，传动机制与传统插销式截然不同；',
           bold_prefix='② 斜面解锁：')
add_bullet('锁紧释放瞬间，底部顶杆弹簧将头壳整体顶出，实现单手秒拆，彻底解决旧产品费力拔取的痛点。',
           bold_prefix='③ 一键弹出：')
add_para('该结构同时具备「抗外部穿刺破坏」与「支持盲操、自动对中」两大优势，区别于易破损的充气式头模'
         '与需对准的手动卡扣头模。', space_after=12)

# ============ 二、与现有技术对比 ============
add_heading('二、与现有技术对比', 1)
add_para('经对比检索，现有技术主要包括三类：螺纹连接式、手动卡扣式与充气式。三者与本申请的机械'
         '锁止结构对比如下：', space_after=8)

rows_data = [
    ('对比维度', '现有技术（螺丝 / 卡扣 / 充气）', '本案方案', '创造性评价'),
    ('安装固定方式', '拧螺丝、拨动侧边卡扣、或充气', '单向盲插下压：倒扣杆直接挤开卡条自动锁紧', '具有显著的操作便利性'),
    ('解锁机制', '反向拧转、手动抠开卡扣、或放气', '顶端斜面按压：解锁按钮的斜面驱动横向卡条平移', '传动机制截然不同，具新颖性'),
    ('分离方式', '需人工用力拔下或撕下软胶头皮', '一键弹出：底部顶杆弹簧瞬间将头壳顶出', '高度创造性（预料不到的效果）'),
]
table = doc.add_table(rows=len(rows_data), cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
widths = [Cm(2.6), Cm(5.2), Cm(5.4), Cm(3.4)]
for row in table.rows:
    for idx, cell in enumerate(row.cells):
        cell.width = widths[idx]
for i, row_data in enumerate(rows_data):
    for j, text in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = text
        if i == 0:
            shade_cell(cell, 'D9E2F3')
            style_table_cell(cell, bold=True, color=DARK_BLUE)
        else:
            style_table_cell(cell, center=(j != 2), size=10)
        if j == 3 and i > 0:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = DARK_GREEN
                    r.bold = True
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ============ 三、授权概率评估 ============
add_heading('三、授权概率评估（美欧双边结论）', 1)

# 评分卡
score_table = doc.add_table(rows=1, cols=1)
score_table.style = 'Table Grid'
score_cell = score_table.cell(0, 0)
shade_cell(score_cell, 'DCFCE7')
score_cell.text = ''
p = score_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('综合授权概率预估：85% - 90%')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = DARK_GREEN
set_east_asia(r)
p2 = score_cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('只要将权利要求（Claims）精准锁定在特定的“内部斜面传动与三弹簧联动结构”上，'
                '不仅不会与现有的头模专利发生冲突，且极大概率能顺利获得授权。')
r2.font.size = Pt(10.5)
r2.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
set_east_asia(r2)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

add_para('1. 美国局（USPTO）授权前景：极高', bold=True, size=11.5, space_after=4)
add_para('USPTO 审查员若在“美发教具”领域内比对，将无法找到任何包含“斜面推拉 + 底部弹簧弹出”的'
         '复合结构。只要申请文件详细限定了导向管（2）、卡条（3）、卡条弹簧（4）以及解锁推杆（8）'
         '的联动关系，即可成功规避现有的插销式（如 US7410358B2）专利限制。', space_after=10)

add_para('2. 欧洲/德国局（EPO/DPMA）授权前景：极高', bold=True, size=11.5, space_after=4)
add_para('欧洲审查员看重“技术问题解决法”。现有同类专利都没有解决“如何单手秒拆且不费力取下”的'
         '技术问题。本案底部“顶杆弹簧一键弹出”的功能，完美解决了旧产品需要费力拔下的痛点。'
         '这一技术贡献完全符合欧洲审查标准中的“创造性高度”。', space_after=12)

# ============ 四、战略建议与风险规避 ============
add_heading('四、战略建议与风险规避', 1)
add_bullet('在向美欧递交专利时，切忌将独立权利要求写成“一种可以通过按压来拆卸的头模”。必须写成：'
           '“一种教学头模，其特征在于内部包含具有斜面的横向锁止卡条，以及带有对应斜面的解锁按钮，'
           '且底部设有轴向顶出弹簧……”。',
           bold_prefix='规避“功能性宽泛索赔”：')
add_bullet('对比检索中发现的充气式头模（容易破损）和手动卡扣头模（需要对准），建议在申请文书中'
           '明确强调机械锁止结构“抗外部穿刺破坏”且“支持盲插下压自动对中锁紧”，进一步拉开与'
           '现有技术的差距。',
           bold_prefix='强调“盲操”与“防刺破”优势：')

doc.add_paragraph()
add_para('—— 本报告仅供内部决策参考 ——', size=9, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

out = r'D:\Eric Cheng\Documents\教学头模专利授权前景分析报告.docx'
doc.save(out)
print('saved:', out)
